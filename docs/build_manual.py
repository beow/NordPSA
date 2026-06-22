"""Genererar docs/NordPSA_Manual.docx från grunden.

Reproducerbar manual-build: kör om efter att flaggor/scenarier/kostnader ändrats.
Källa till sanning för flaggorna = scripts/run_model.py (argparse). Uppdatera
texterna här när modellen ändras. Kör:  python docs/build_manual.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = Path(__file__).with_name("NordPSA_Manual.docx")

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY   = RGBColor(0x55, 0x55, 0x55)
MONO   = "Consolas"

doc = Document()

# ---- bas-stilar ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

def _set_heading_color(style_name, color, size):
    st = doc.styles[style_name]
    st.font.color.rgb = color
    st.font.size = Pt(size)
    st.font.name = "Calibri"

for nm, sz in [("Title", 24), ("Heading 1", 16), ("Heading 2", 12.5), ("Heading 3", 11)]:
    _set_heading_color(nm, ACCENT, sz)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)

def para(t="", italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.italic = italic
    if color: r.font.color.rgb = color
    if size:  r.font.size = Pt(size)
    return p

def bullet(t, sub=False):
    p = doc.add_paragraph(style="List Bullet 2" if sub else "List Bullet")
    # stöd för **fetstil**-markörer
    parts = t.split("**")
    for i, seg in enumerate(parts):
        run = p.add_run(seg)
        if i % 2 == 1:
            run.bold = True
    return p

def code(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    for i, ln in enumerate(lines):
        r = p.add_run(("\n" if i else "") + ln)
        r.font.name = MONO
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, hd in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        run = c.paragraphs[0].add_run(hd)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if str(val).startswith("--") or str(val).startswith("ZON"):
                run.font.name = MONO
                run.font.size = Pt(8.5)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    doc.add_paragraph()
    return t

# =====================================================================
# TITEL
# =====================================================================
ti = doc.add_paragraph(style="Title")
ti.add_run("NordPSA")
sub = doc.add_paragraph()
sub.add_run("Nordisk kraftsystemmodell — modellbeskrivning och användarmanual").italic = True
meta = doc.add_paragraph()
mr = meta.add_run("PyPSA-baserad LP-dispatch + kapacitetsexpansion · 6 zoner · 2023–2025\n"
                  "Version juni 2026 (reviderad — sektorkoppling, SvK-scenarier)")
mr.font.color.rgb = GREY
mr.font.size = Pt(9.5)

# =====================================================================
# INNEHÅLL
# =====================================================================
h1("Innehåll")
toc = [
    "1.  Vad är NordPSA",
    "2.  Arkitektur och komponenter",
    "3.  Installation och uppsättning",
    "4.  Datapipeline",
    "5.  Köra modellen",
    "6.  Scenarier (kostnad + efterfrågan)",
    "7.  Sektorkoppling (värme, EV, vätgas)",
    "8.  Kostnadsmodellen",
    "9.  Funktionsreferens (flaggor)",
    "10. Solver och prestanda",
    "11. Genomgångna exempel",
    "12. Konfiguration (zones.yaml)",
    "13. Bilaga: fullständig flaggtabell",
]
for t in toc:
    para(t)

# =====================================================================
# 1. VAD ÄR NORDPSA
# =====================================================================
h1("1.  Vad är NordPSA")
para("NordPSA är en modell av det nordiska kraftsystemet byggd på PyPSA. Den kombinerar "
     "linjär dispatch-optimering (LP) med kapacitetsexpansion (investering) för sex "
     "aggregerade zoner och täcker 2023–2025 i tim- eller 3-timmarsupplösning. Modellen "
     "har vuxit från ren dispatch till en sektorkopplad expansionsmodell med värme, "
     "vägtransport (EV), vätgas och SvK-baserade 2040-scenarier.")
para("Zonerna är:")
table(["Zon", "Område", "Karaktär"], [
    ["SE-N", "SE1+SE2", "Vattenkraftsrik, exporterande, ingen kontinentkoppling"],
    ["SE-S", "SE3+SE4", "Kärnkraft + last, kontinentkopplad (DE/PL/LT)"],
    ["NO-N", "NO3+NO4", "Vattenkraft, ~30 TWh reservoar, ingen kontinentkoppling"],
    ["NO-S", "NO1+NO2+NO5", "Stor vattenkraft, ~57 TWh reservoar, kontinentkopplad"],
    ["DK", "DK1+DK2", "Vindtung, djup kontinentkoppling (DE/NL/GB)"],
    ["FI", "FI", "Kärnkraft + import (EE), måttlig vattenkraft"],
], widths=[0.7, 1.3, 4.2])
para("Modellen minimerar systemets totala kostnad (CAPEX + drift) för timvis dispatch och "
     "valfri investering, givet last, väder (VRE-profiler), hydrologi och prissatta "
     "marknadsanslutningar mot kontinenten.")

# =====================================================================
# 2. ARKITEKTUR
# =====================================================================
h1("2.  Arkitektur och komponenter")
h2("2.1  Nätverkskomponenter (PyPSA)")
table(["Komponent", "Carrier", "Beskrivning"], [
    ["Bus", "AC", "En elbuss per zon"],
    ["Bus", "H2 / heat / EV / chp fuel", "Sektorbussar per zon (tillval, se kap. 7)"],
    ["Link", "—", "Bidirektionell NTC mellan zoner (p_min_pu=−1)"],
    ["Link", "electrolyser / heat hp / chp …", "Omvandlingslänkar el→H2/värme (sektorkoppling)"],
    ["StorageUnit", "hydro", "Reservoar med faktiskt inflöde, cyklisk SOC, ingen pumpning"],
    ["StorageUnit", "battery", "Li-ion, cyklisk SOC (tillval)"],
    ["Generator", "nuclear", "Befintlig flotta fast; ny via --add-nuclear (syntetisk, extendable)"],
    ["Generator", "wind_onshore/offshore, solar", "VRE med kapacitetsfaktor-profiler, extendable"],
    ["Generator", "thermal", "Must-run (p_min=p_max=faktisk profil)"],
    ["Generator", "gas", "Dispatchbar peaker, extendable"],
    ["Generator", "market", "Import/export-ventil mot kontinentpris (pris-elastisk trappa)"],
    ["Generator", "slack", "Lastavkoppling (3000 EUR/MWh)"],
    ["Store / Load", "H2 / heat / EV", "Lager + baslast i sektorbussarna (tillval)"],
], widths=[1.3, 1.8, 3.1])

h2("2.2  Dataflöde")
code([
    "eSett API     → data/raw/production_*.parquet   (last + produktion per carrier)",
    "ENTSO-E       → reservoar + produktion (hydro), NO-priser",
    "NVE           → norska reservoarnivåer (inflöde)",
    "Energy Charts → VRE-profiler + DE-LU dagspris",
    "Renewables.ninja → syntetiska offshore-profiler (kap. 7)",
    "        │",
    "        ▼",
    "scripts/build_inputs.py  →  data/processed/*.parquet + *.yaml",
    "        │",
    "        ▼",
    "nordpsa/network.py  →  pypsa.Network",
    "        │",
    "        ▼",
    "scripts/run_model.py  →  results/<output>/",
])

h2("2.3  Paketstruktur")
for t in [
    "**nordpsa/network.py** — bygger PyPSA-nätet från processad indata (build_network).",
    "**nordpsa/esett.py** — eSett open-data-klient, aggregerar MBA:er till zoner.",
    "**nordpsa/ec.py** — Energy Charts-klient (VRE-profiler, DE-LU-pris).",
    "**nordpsa/entsoe.py** — ENTSO-E (reservoar, produktion, NO-priser).",
    "**nordpsa/ninja.py** — Renewables.ninja-klient (offshore-profiler).",
    "**nordpsa/hydro.py** — inflödesmodell (faktiska NVE/ENTSO-E-profiler + parametrisk fallback).",
    "**scripts/run_model.py** — CLI: bygger och löser modellen.",
    "**scripts/build_inputs.py / build_heat.py** — bygger processad indata resp. värmeprofiler.",
]:
    bullet(t)

# =====================================================================
# 3. INSTALLATION
# =====================================================================
h1("3.  Installation och uppsättning")
code([
    "conda env create -f environment.yml",
    "conda activate spyder-env        # arbetsmiljön",
    "pip install -e .                 # installera nordpsa-paketet",
])
para("Obs: den fungerande conda-miljön heter spyder-env.", italic=True, color=GREY)

# =====================================================================
# 4. DATAPIPELINE
# =====================================================================
h1("4.  Datapipeline")
para("Data måste hämtas och byggas innan modellen kan lösas:")
code([
    "make fetch      # eSett last + produktion",
    "make fetch-ec   # Energy Charts VRE-profiler + DE-LU-pris",
    "make build      # bygg processade parquet-inputs",
    "make solve      # kör modellen",
])
para("Processad indata (data/processed/): load.parquet, vre_profiles.parquet, vre_pnom.yaml, "
     "nuclear_profile.parquet, thermal_profile.parquet, market_prices.parquet, samt "
     "heat_load.parquet (värme, byggs av build_heat.py). Hydrologiska inflödesprofiler "
     "hämtas från NVE+ENTSO-E (faktiska data, default).")
para("Hydro-parametrar i config/hydro_params.yaml är manuellt kalibrerade och regenereras "
     "INTE av build_inputs.py (skilj från data/processed/hydro_params.yaml som är auto-genererad "
     "och inte ska användas).", italic=True, color=GREY)

# =====================================================================
# 5. KÖRA MODELLEN
# =====================================================================
h1("5.  Köra modellen")
h2("5.1  Grundläggande")
code([
    "# Ren dispatch (3h), hela perioden",
    "python scripts/run_model.py --resolution 3 --no-expansion --output run_baseline",
    "",
    "# Kapacitetsexpansion (kom ihåg --spill-cost 50!)",
    "python scripts/run_model.py --resolution 3 --spill-cost 50 --output run_expansion",
    "",
    "# Ett enstaka år, timupplösning",
    "python scripts/run_model.py --resolution 1 --year 2024 --no-expansion --output run_2024",
])
h2("5.2  Resultat")
para("Resultat hamnar i results/<output>/ (gitignorerat):")
for t in [
    "**network.nc** — fullständigt PyPSA-nät (inkl. inflöde, dispatch, p_nom_opt).",
    "**prices.csv** — buspriser (EUR/MWh) per zon och timme.",
    "**dispatch_generators.csv, dispatch_hydro.csv, flows.csv** — dispatch och flöden.",
    "**hydro_soc.csv, hydro_spill.csv, water_value.csv** — hydrologi och vattenvärde.",
    "**h2_store_soc.csv** — vätgaslager-SOC (om H2 finns).",
    "**run_meta.txt** — körningens syfte (--desc), flaggor, git-state, tid.",
]:
    bullet(t)
h2("5.3  Körningsdisciplin")
para("Committa EFTER en lyckad simulering, inte före — så att bara goda körningar spåras till "
     "kodtillstånd. Namnge resultatmappen runXXX_<beskrivning> med löpnummer och ange den i "
     "commit-meddelandet för spårbarhet. Använd --dry-run för att verifiera nätbygget innan en "
     "tung körning startas.")

# =====================================================================
# 6. SCENARIER
# =====================================================================
h1("6.  Scenarier (kostnad + efterfrågan)")
para("Två ortogonala scenariolager läses ur config/zones.yaml och kombineras fritt. De är "
     "avsedda för expansionskörningar (framtidsår).")
h2("6.1  Kostnadsscenario  (--cost-scenario)")
para("Skriver över costs-blocket med en framtida teknikkostnadsuppsättning ur cost_scenarios. "
     "Inkluderar byggränta (IDC, ränta under byggtid): overnight' = overnight·(1 + build_years/2·r). "
     "Tillgängliga: svk_2040, svk_2050.")
h2("6.2  Efterfrågescenario  (--demand-scenario)")
para("Adderar ett efterfrågetillskott ADDITIVT ovanpå eSett-basen (basen bevaras). Per zon: "
     "extra flat last, vätgaslast (H2), elbilar (EV), VRE-utbyggnadstak (p_nom_max), exogena "
     "golv (p_nom_min) och NTC-höjningar. Tillgängligt: svk_2040_mm (SvK Långsiktig "
     "marknadsanalys 2026, scenario Mot Mål).")
para("svk_2040_mm i korthet: per-land landstotaler 2040 (TWh) SE 243 / NO 194 / FI 145 / DK 81. "
     "Landstotaler + SE:s H2 (24 TWh) är hårda SvK-tal; zon-split inom land, FI/DK:s H2-fördelning "
     "och VRE-potentialernas zonandelar är FLAGGADE ANTAGANDEN. NO är kalibrerat mot Statnett "
     "(ej SvK-split). VRE-potentialer (tak/golv) från LMA2026 Bilaga B, omräknade TWh→MW vid "
     "antagen CF.", color=GREY, size=9.5)
para("CLI-givna --add-h2/--add-ev har företräde och skrivs inte över av scenariot.", italic=True, color=GREY)

# =====================================================================
# 7. SEKTORKOPPLING
# =====================================================================
h1("7.  Sektorkoppling (värme, EV, vätgas)")
para("Sektorerna är opt-in och adderar bussar + omvandlingslänkar per zon. Dagens sektor-el "
     "dras bort ur AC-lasten där relevant (undviker dubbelräkning).")

h2("7.1  Fjärrvärme  (--add-heat)")
para("Per-zon värmebuss (SE-N/SE-S/DK/FI) med: FV-behovsprofil (When2Heat+Open-Meteo), "
     "ackumulator (termiskt lager, e_cyclic), el-panna, stor-VP (COP per zon), bio-grundlast och "
     "bakpress-KVV (multi-output Link bränsle→el+värme). KVV-elen blir endogen och must-run-"
     "termiken reduceras med share_of_thermal (ingen dubbelräkning). Värmelagret är fast per zon "
     "(store_gwh) eller investerbart med --heat-store-ext.")
para("--chp-fixed-gw GW är ett snabbalternativ: kopplar bort hela värmebussen och återinför bara "
     "KVV-elen som ett fast must-run-block (totala toppen = GW). Ger ~samma systembild som full "
     "värmesektor men betydligt mindre LP (snabbare solve). Ömsesidigt uteslutande med --add-heat.")

h2("7.2  Elfordon  (--add-ev)")
para("SvK-formulering (config ev.mode=svk): förbrukningen delas 50/50 flexibel/oflexibel. "
     "Oflexibel halva = fast AC-last formad som körprofilen (middagsdippen överbryggd). Flexibel "
     "halva = SvK-reservoar (batteri 12 GWh per flexibel-TWh, laddtak = batteri/10h, fri "
     "tidsförflyttning över dygnets billigaste timmar). Äldre per-fordon-läge finns som ev.mode=fleet.")

h2("7.3  Vätgas  (--add-h2 / --add-h2-ext)")
para("Power-to-X per zon: elbuss → Link(elektrolys, η=0.66) → H2-buss → Store(lager, e_cyclic) "
     "+ Load(konstant baslast) + valfri turbin tillbaka till el. --add-h2 ger fasta storlekar; "
     "--add-h2-ext gör elektrolys + lager investerbara. I svk_2040_mm är H2-lasten exogent fast "
     "medan elektrolysören är endogen (dimensioneras). Lagerkostnad är geologiberoende (urberg "
     "default, billig saltkavern för DK).")

# =====================================================================
# 8. KOSTNADSMODELLEN
# =====================================================================
h1("8.  Kostnadsmodellen")
para("Kapitalkostnad annualiseras med Capital Recovery Factor (CRF):")
code([
    "CRF = r·(1+r)^L / ((1+r)^L − 1)        r = 0.06 (kalkylränta), L = livslängd",
    "capital_cost = overnight · (CRF + FOM) · n_years     FOM = 2% (default)",
])
para("Kapitalkostnaden läggs på p_nom_opt (total installerad kapacitet). Alla extendable "
     "generatorer har p_nom_min = befintlig kapacitet. n_years skalar årskostnaden till "
     "körningens horisont så att CAPEX och OPEX väger rätt. Ny kärnkraft kan få egen kalkylränta "
     "per zon med --nuclear-discount-rate (påverkar bara den extendable expansionen).")

h2("8.1  Baskostnader (config costs, default)")
para("Annualiserat = overnight·(CRF+FOM). Vid --cost-scenario svk_2040/svk_2050 byts dessa mot "
     "scenariots värden (inkl. byggränta).", color=GREY, size=9.5)
table(["Teknik", "Overnight", "Livsl. (år)", "VOM €/MWh", "Annual. €/kW/år"], [
    ["nuclear", "7.0 €/W", "60", "20", "573"],
    ["wind_onshore", "1.2 €/W", "30", "1.6", "100"],
    ["wind_offshore", "2.2 €/W", "30", "3.0", "199"],
    ["solar", "0.4 €/W", "35", "1.0", "35"],
    ["gas (peaker)", "0.45 €/W", "30", "100", "51"],
    ["hydro", "ej extendable", "—", "0.6", "—"],
], widths=[1.6, 1.3, 1.1, 1.1, 1.4])

h2("8.2  Vätgas (kalibrerat mot DEA)")
table(["Komponent", "Overnight", "Verkningsgrad", "Annual."], [
    ["Elektrolysör (alkalisk)", "575 €/kW", "0.66 (LHV)", "68 €/kW/år"],
    ["Turbin (H2-OCGT)", "435 €/kW", "0.41 (el)", "47 €/kW/år"],
    ["Lager — urberg (SE/NO/FI)", "5 €/kWh", "—", "0.46 €/kWh/år"],
    ["Lager — saltkavern (DK)", "0.5 €/kWh", "—", "0.046 €/kWh/år"],
], widths=[2.3, 1.3, 1.3, 1.3])

# =====================================================================
# 9. FUNKTIONSREFERENS
# =====================================================================
h1("9.  Funktionsreferens (flaggor)")
para("Grupperad genomgång. Fullständig tabell i kap. 13.", italic=True, color=GREY)

h2("9.1  Dispatch vs expansion")
for t in [
    "**--no-expansion** — lås all kapacitet (ren dispatch).",
    "**--spill-cost 50** — KRÄVS i expansionskörningar (default 0.1 ger phantom spill + VRE-överinvestering).",
    "**--extra-load MW** — extra flat last per zon (stresstest).",
    "**--dry-run** — bygg nätet + komponent-sammanfattning men solve:a inte (verifiering).",
    "**--desc TEXT** — körningens syfte, sparas i run_meta.txt.",
]:
    bullet(t)

h2("9.2  Scenarier")
for t in [
    "**--cost-scenario NAMN** — framtida teknikkostnader (svk_2040, svk_2050), inkl. byggränta.",
    "**--demand-scenario NAMN** — additivt efterfrågetillskott (svk_2040_mm): last, H2, EV, VRE-tak/golv, NTC.",
]:
    bullet(t)

h2("9.3  Hydrologi")
for t in [
    "**--spill-cost EUR** — hydro-spillkostnad (expansion: 50).",
    "**--soc-pin ZON:START:END** — icke-cyklisk: lås båda reservoar-ändpunkterna till faktiska fyllnadsfraktioner (kalibrering).",
]:
    bullet(t)
para("Vattenvärdet = dualvariabeln på reservoarens energibalans (mu_energy_balance), prisgolvet. "
     "Sparas till water_value.csv.", color=GREY, size=9.5)

h2("9.4  Kärnkraft, vind, sol")
for t in [
    "**--add-nuclear ZON:N:SEED** — N nya reaktorer, syntetisk stokastisk tillgänglighet (seed), EXTENDABLE (tak N×1500 MW). Trigger: befintlig flotta byter till syntetisk profil.",
    "**--nuclear-discount-rate ZON:RATE** — egen kalkylränta för NY kärnkraft per zon (t.ex. SE-N:0.03).",
    "**--add-wind ZON:MW** — fast landvind (dispatch, ej extendable); energi-motsvarighet till --add-nuclear.",
    "**--onwind-capfac-increase FRAC** — höj landvinds CF (olinjär potens-transform, märkeffekt oförändrad).",
    "**--offwind-capfac-increase FRAC** — som ovan för havsvind (2040-nybyggnadsflotta).",
    "**--onshore-cap ZON:MW** — expansionstak landvind per zon.",
    "**--solar-cap MW** — expansionstak sol i alla zoner.",
]:
    bullet(t)

h2("9.5  Lagring")
for t in [
    "**--add-battery ZON:MW:HOURS** — fast batteri; **--battery-extendable** gör det investerbart.",
    "**--battery 4h [ZON:MW ...]** — expanderbart batteri i alla zoner (utan zonlista) eller fasta storlekar (med).",
]:
    bullet(t)

h2("9.6  Sektorkoppling")
for t in [
    "**--add-heat** — fjärrvärmesektor (heat-buss + ackumulator + el-panna + VP + bio/KVV).",
    "**--heat-store-ext** — gör värmeackumulatorn investerbar (kräver --add-heat).",
    "**--chp-fixed-gw GW** — snabbalternativ: KVV-el som fast must-run-block, ingen värmebuss.",
    "**--add-ev ZON:N_CARS:N_HEAVY** — fordonsladdning (SvK flex/oflex-formulering).",
    "**--add-h2 ZON:DEMAND:EL:STORE[:TURB]** — vätgassystem (fasta storlekar).",
    "**--add-h2-ext ZON:DEMAND[:STORE_MAX]** — investerbart vätgassystem (elektrolys + lager).",
]:
    bullet(t)

h2("9.7  Riktad budgetbegränsad VRE-expansion")
for t in [
    "**--expand-vre ZON** — gör bara den zonens vind+sol investerbara (oavsett --no-expansion).",
    "**--expand-budget-musd/-meur** — tak på overnight-kostnaden → tvingad likhet + kapital=0 (ren dispatcheffekt).",
]:
    bullet(t)

h2("9.8  Marknad / NTC")
for t in [
    "**--market-elasticity** (default PÅ) — pris-elastisk kontinentgräns (trappa): stora flöden flyttar gränspriset. **--no-market-elast** stänger av.",
    "**--continent-diurnal-scale FAKTOR** — komprimera kontinentventilens dygnssväng (1.0=oförändrat, 0.5=halverad).",
    "**--effective-ntc** — effektiv kontinentkapacitet (P80) istället för märkkapacitet.",
    "**--market-scale FACTOR|ZON:F,...** — skala kontinentkablar.",
    "**--no-market** — stäng alla kontinentanslutningar.",
    "**--voll** — VOLL-slack (3000 €/MWh) i alla zoner (förhindrar dualexplosion).",
]:
    bullet(t)

# =====================================================================
# 10. SOLVER
# =====================================================================
h1("10.  Solver och prestanda")
para("HiGHS interior point (IPM/IPX) + crossover. Crossover KRÄVS för korrekt "
     "kapacitetsexpansion — utan den ger IPM en inre (utsmetad) lösning där p_nom_opt fastnar "
     "nära p_nom_min. Konfigureras i zones.yaml (solver-blocket; ipm_optimality_tolerance 1e-6).")
for t in [
    "**IPM (interior point)** — går genom polytopens inre; få men dyra iterationer; skalar bra för stora glesa LP.",
    "**Simplex** — går längs kanter hörn-till-hörn; ger exakt baslösning; bra för varmstart, känslig för degeneration.",
    "**Crossover** — trycker IPM-lösningen till ett hörn → exakta p_nom_opt. 'Imprecise' = giltig men städas med simplex.",
]:
    bullet(t)
para("Prestanda-tumregler: expansion i 3h ≈ 1–2 h solve; full 2h-körning ≈ 4–5 h (håller "
     "minnet med KVV-fast); 1h-expansion riskerar OOM i extraktionen. Degeneration (t.ex. när "
     "billig kärnkraft blir kostnadsekvivalent med VRE) kan stalla IPM eller ge divergerande "
     "simplex-städning — botas av hårda kapacitetstak och brutna kostnads-likheter.")

# =====================================================================
# 11. EXEMPEL
# =====================================================================
h1("11.  Genomgångna exempel")

h2("11.1  SvK 2040-expansion (sektorkopplad)")
para("Full framtidskörning: 2040-kostnader + additiv efterfrågan + värme + EV + ny kärnkraft.")
code([
    "python scripts/run_model.py --resolution 3 --spill-cost 50 --market-elasticity \\",
    "  --cost-scenario svk_2040 --demand-scenario svk_2040_mm \\",
    "  --add-heat --add-nuclear SE-S:10:201 SE-N:10:202 FI:10:203 \\",
    "  --onwind-capfac-increase 0.30 --offwind-capfac-increase 0.10 \\",
    "  --output run_svk2040",
])

h2("11.2  Snabb variant med fast KVV (mindre LP)")
code([
    "python scripts/run_model.py --resolution 3 --spill-cost 50 --market-elasticity \\",
    "  --cost-scenario svk_2040 --demand-scenario svk_2040_mm \\",
    "  --chp-fixed-gw 5 --add-nuclear SE-S:10:201 \\",
    "  --output run_chpfixed",
])

h2("11.3  VRE+batteri vs kärnkraft för samma budget")
para("Vad gör samma 'tilläggsnota' med priserna om den läggs på VRE+batteri istället för "
     "kärnkraft i SE-S? (tvingad likhet → ren dispatcheffekt)")
code([
    "python scripts/run_model.py --resolution 1 --no-expansion \\",
    "  --expand-vre SE-S --expand-budget-meur 20000 \\",
    "  --add-battery SE-S:50000:4 --battery-extendable --output run_vre_budget",
])
para("Slutsats (tidigare körningar): VRE+batteri sänker medelpriset mest per krona, med tydliga "
     "diminishing returns. En must-run-enhet är aldrig marginell (sätter aldrig priset); en "
     "dispatchbar enhet sätter priset till sin MC när den är interiör.", color=GREY, size=9.5)

# =====================================================================
# 12. KONFIGURATION
# =====================================================================
h1("12.  Konfiguration (zones.yaml)")
para("All parametrisering ligger i config/zones.yaml:")
for t in [
    "**zones** — zondefinitioner med hydro/kärnkraft befintlig kapacitet, reservoarstorlek, SOC-startnivå.",
    "**links** — interna NTC mellan zoner.",
    "**market_connections** (+ _effective_mw) — kontinentanslutningar och priszon.",
    "**market_elasticity** — trappa för pris-elastisk kontinentgräns.",
    "**costs** — baskostnader (overnight, livslängd, VOM, extendable, p_nom_max), inkl. battery och hydrogen.",
    "**cost_scenarios / demand_scenarios** — framtidsscenarier (svk_2040/2050, svk_2040_mm).",
    "**heat / ev** — sektorparametrar (värmezoner, COP, KVV; EV-läge och fordonstal).",
    "**hydrogen** — valfria vätgassystem per zon (opt-in; default avstängt).",
    "**solver** — HiGHS IPM + crossover, toleranser.",
    "**snapshots** — simuleringsperiod och upplösning.",
]:
    bullet(t)
para("Viktiga designval: IPM med crossover krävs för korrekt expansion; termisk och befintlig "
     "kärnkraft är must-run; hydro-SOC är cyklisk (start=slut); vätgaslagrets kostnad är per zon "
     "efter geologi; efterfrågescenarier är additiva över eSett-basen.")

# =====================================================================
# 13. BILAGA: FLAGGTABELL
# =====================================================================
h1("13.  Bilaga: fullständig flaggtabell")
table(["Flagga", "Funktion"], [
    ["--resolution H", "Tidsupplösning i timmar (1, 2 eller 3)"],
    ["--year Y", "Kör ett enstaka år"],
    ["--output NAMN", "Resultatmapp under results/"],
    ["--desc TEXT", "Körningens syfte → run_meta.txt"],
    ["--dry-run", "Bygg nät + sammanfattning, ingen solve"],
    ["--no-expansion", "Ren dispatch (ingen investering)"],
    ["--spill-cost EUR", "Hydro-spillkostnad (expansion: 50)"],
    ["--extra-load MW", "Extra flat last per zon"],
    ["--cost-scenario NAMN", "Framtida teknikkostnader (svk_2040/2050)"],
    ["--demand-scenario NAMN", "Additivt efterfrågetillskott (svk_2040_mm)"],
    ["--soc-pin ZON:START:END", "Lås reservoar-ändpunkter (kalibrering)"],
    ["--add-nuclear ZON:N:SEED", "Ny syntetisk kärnkraft (extendable)"],
    ["--nuclear-discount-rate ZON:RATE", "Egen kalkylränta för ny kärnkraft"],
    ["--add-wind ZON:MW", "Fast landvind (dispatch)"],
    ["--onwind-capfac-increase FRAC", "Höj landvinds CF"],
    ["--offwind-capfac-increase FRAC", "Höj havsvinds CF"],
    ["--onshore-cap ZON:MW", "Expansionstak landvind per zon"],
    ["--solar-cap MW", "Expansionstak sol (alla zoner)"],
    ["--add-battery ZON:MW:HOURS", "Fast batteri"],
    ["--battery-extendable", "Gör --add-battery investerbart"],
    ["--battery 4h [ZON:MW ...]", "Expanderbart/fast batteri (varaktighet fast)"],
    ["--add-heat", "Fjärrvärmesektor"],
    ["--heat-store-ext", "Investerbart värmelager (kräver --add-heat)"],
    ["--chp-fixed-gw GW", "Fast KVV-block, ingen värmebuss"],
    ["--add-ev ZON:CARS:HEAVY", "Fordonsladdning (SvK flex/oflex)"],
    ["--add-h2 ZON:DEMAND:EL:STORE[:TURB]", "Vätgassystem (fast storlek)"],
    ["--add-h2-ext ZON:DEMAND[:STORE_MAX]", "Investerbart vätgassystem"],
    ["--expand-vre ZON", "Riktad VRE-expansion i zon"],
    ["--expand-budget-musd/-meur", "Overnight-budgettak (tvingad likhet)"],
    ["--market-elasticity / --no-market-elast", "Pris-elastisk kontinentgräns (default på)"],
    ["--continent-diurnal-scale FAKTOR", "Komprimera kontinentventilens dygnssväng"],
    ["--effective-ntc", "P80 kontinentkapacitet"],
    ["--market-scale FACTOR|ZON:F,...", "Skala kontinentkablar"],
    ["--no-market", "Stäng kontinentanslutningar"],
    ["--voll", "VOLL-slack i alla zoner"],
], widths=[2.7, 3.7])

doc.save(OUT)
print(f"Sparad: {OUT}  ({OUT.stat().st_size} bytes)")
